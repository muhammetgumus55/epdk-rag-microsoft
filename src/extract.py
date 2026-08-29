"""Extract text from the EPDK mevzuat corpus (.doc / .docx / .pdf).

Design notes grounded in what this corpus actually is:

* EPDK filenames are opaque and derived from download IDs, so a file's name is
  never treated as its identity -- it is carried only as `original_filename`
  metadata. Identity comes from content (see src.titles).
* File type is decided by magic bytes, not the extension. The corpus contains
  at least one .pdf that is really a .docx, and ZIP-magic covers .docx, .xlsx
  and plain archives alike, so the container is opened to tell them apart.
* .doc is converted to .docx (not to .txt) via LibreOffice. Converting to text
  flattens tables away, and LibreOffice's default text filter writes cp1254 on
  a Turkish Windows box -- decoding that as UTF-8 raises, and decoding it as
  cp1252 silently yields plausible-but-wrong characters. Going through .docx
  keeps table structure and stays UTF-8 end to end.
* Paragraphs and tables are read in document order by walking the body XML.
  Some corpus documents put ALL their content in tables, so paragraph-only
  extraction would report them as empty.
"""
from __future__ import annotations

import hashlib
import shutil
import subprocess
import sys
import tempfile
import unicodedata
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from . import config

# --------------------------------------------------------------------------
# Turkish-aware case folding
# --------------------------------------------------------------------------
# Python's default casing corrupts Turkish: "İSTANBUL".lower() yields
# "i̇stanbul" (i + U+0307 combining dot), and "ırmak".upper() yields "IRMAK"
# which then lowercases back to "irmak", losing the dotless i. Every comparison
# or normalization of Turkish text in this package must use these helpers.

_TR_LOWER_MAP = str.maketrans({"I": "ı", "İ": "i"})
_TR_UPPER_MAP = str.maketrans({"ı": "I", "i": "İ"})

TURKISH_CHARS = "ıİşŞğĞüÜöÖçÇ"


def tr_lower(text: str) -> str:
    """Lowercase using Turkish rules (I->ı, İ->i)."""
    return text.translate(_TR_LOWER_MAP).lower()


def tr_upper(text: str) -> str:
    """Uppercase using Turkish rules (ı->I, i->İ)."""
    return text.translate(_TR_UPPER_MAP).upper()


# --------------------------------------------------------------------------
# Data model
# --------------------------------------------------------------------------


@dataclass
class Page:
    """One page of extracted text. `number` is 1-based, or None when unknown."""

    number: int | None
    text: str
    needs_ocr: bool = False


@dataclass
class ExtractedDoc:
    """One successfully extracted document: its pages, identity, and quality flags."""

    path: Path
    original_filename: str
    doc_id: str
    file_sha256: str
    detected_type: str
    pages: list[Page] = field(default_factory=list)
    flags: list[str] = field(default_factory=list)
    # Why page numbers are absent, when they are. None means pages are real.
    page_number_note: str | None = None

    @property
    def text(self) -> str:
        """The whole document as one string, pages joined and blanks dropped."""
        return "\n\n".join(p.text for p in self.pages if p.text.strip())

    @property
    def has_pages(self) -> bool:
        """Whether this document has a real page model (PDF) vs none (docx)."""
        return any(p.number is not None for p in self.pages)


@dataclass
class SkippedDoc:
    """A file that was not extracted, with the reason why."""

    path: Path
    reason: str


class ExtractionError(Exception):
    """Raised when a document cannot be extracted at all."""


class LibreOfficeUnavailable(Exception):
    """Raised when .doc files are present but LibreOffice cannot be used."""


# --------------------------------------------------------------------------
# Content-based type detection
# --------------------------------------------------------------------------

# Types this module can turn into text. Anything else is reported, not dropped.
EXTRACTABLE = {"doc", "docx", "pdf"}


def detect_type(path: Path) -> str:
    """Identify a file by magic bytes, opening ZIP containers to disambiguate.

    Returns one of: doc, docx, pdf, xlsx, pptx, xls, zip, rtf, html, empty,
    or 'unknown:<hex>'.
    """
    try:
        with open(path, "rb") as fh:
            head = fh.read(8)
    except OSError as exc:
        raise ExtractionError(f"cannot read file: {exc}") from exc

    if not head:
        return "empty"
    if head.startswith(b"%PDF"):
        return "pdf"
    if head.startswith(b"{\\rtf"):
        return "rtf"
    if head.startswith(b"\xd0\xcf\x11\xe0"):
        # OLE2 compound file: legacy .doc, .xls or .ppt.
        return _detect_ole2(path)
    if head.startswith(b"PK\x03\x04"):
        return _detect_zip(path)
    if tr_lower(head.decode("latin-1", "replace")).startswith(("<html", "<!doc")):
        return "html"
    return "unknown:" + head[:4].hex()


# OLE2 stream names live in the directory sectors, whose offset varies with the
# file's layout -- a corpus .xls was found with its "Workbook" entry well past
# the 8 KB mark, so probing only the header misclassifies it as .doc. Scan the
# whole file, capped so a pathologically large one cannot blow up memory.
_OLE2_PROBE_BYTES = 8 * 1024 * 1024


def _detect_ole2(path: Path) -> str:
    """Distinguish Word/Excel/PowerPoint inside an OLE2 container.

    Stream names appear as UTF-16 in the raw bytes, so a substring probe avoids
    pulling in an OLE parsing dependency.
    """
    try:
        with open(path, "rb") as fh:
            blob = fh.read(_OLE2_PROBE_BYTES)
    except OSError:
        return "doc"
    if b"W\x00o\x00r\x00d\x00D\x00o\x00c\x00u\x00m\x00e\x00n\x00t" in blob:
        return "doc"
    if b"W\x00o\x00r\x00k\x00b\x00o\x00o\x00k" in blob or b"B\x00o\x00o\x00k" in blob:
        return "xls"
    if b"P\x00o\x00w\x00e\x00r\x00P\x00o\x00i\x00n\x00t" in blob:
        return "ppt"
    return "doc"  # OLE2 in this corpus is overwhelmingly Word


def _detect_zip(path: Path) -> str:
    """Open a ZIP container to tell .docx from .xlsx from a plain archive."""
    try:
        with zipfile.ZipFile(path) as zf:
            names = zf.namelist()
    except (zipfile.BadZipFile, OSError):
        return "zip"
    if any(n.startswith("word/") for n in names):
        return "docx"
    if any(n.startswith("xl/") for n in names):
        return "xlsx"
    if any(n.startswith("ppt/") for n in names):
        return "pptx"
    return "zip"


def sha256_file(path: Path) -> str:
    """Hash the SOURCE FILE (not extracted text) so re-runs can skip unchanged files."""
    h = hashlib.sha256()
    with open(path, "rb") as fh:
        for block in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def make_doc_id(path: Path, file_hash: str) -> str:
    """Stable internal identifier, deliberately independent of the opaque filename.

    The hash prefix is the identity; the slug is a human-readable convenience
    only, and must never be parsed for meaning.
    """
    stem = unicodedata.normalize("NFKD", path.stem)
    stem = stem.encode("ascii", "ignore").decode("ascii")
    slug = "".join(ch if ch.isalnum() else "-" for ch in tr_lower(stem))
    slug = "-".join(part for part in slug.split("-") if part)[:48].strip("-")
    return f"{file_hash[:12]}-{slug}" if slug else file_hash[:12]


# --------------------------------------------------------------------------
# LibreOffice (.doc -> .docx)
# --------------------------------------------------------------------------

_WINDOWS_SOFFICE = [
    Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
    Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
]


def find_libreoffice() -> Path | None:
    """Locate the soffice binary, or None."""
    if config.LIBREOFFICE_PATH:
        candidate = Path(config.LIBREOFFICE_PATH)
        return candidate if candidate.exists() else None
    for name in ("soffice", "soffice.exe", "libreoffice"):
        found = shutil.which(name)
        if found:
            return Path(found)
    for candidate in _WINDOWS_SOFFICE:
        if candidate.exists():
            return candidate
    return None


def convert_doc_batch(paths: list[Path], soffice: Path, outdir: Path) -> dict[Path, Path]:
    """Convert .doc files to .docx in one soffice invocation.

    Returns {source: converted}. Missing keys mean that file failed to convert;
    the caller reports those rather than dropping them silently.
    """
    if not paths:
        return {}
    outdir.mkdir(parents=True, exist_ok=True)
    cmd = [
        str(soffice),
        "--headless",
        "--norestore",
        "--convert-to",
        "docx:MS Word 2007 XML",
        "--outdir",
        str(outdir),
        *(str(p.resolve()) for p in paths),
    ]
    try:
        subprocess.run(
            cmd,
            check=False,
            capture_output=True,
            timeout=config.LIBREOFFICE_TIMEOUT,
        )
    except subprocess.TimeoutExpired:
        return {}

    result: dict[Path, Path] = {}
    for src in paths:
        produced = outdir / (src.stem + ".docx")
        if produced.exists() and produced.stat().st_size > 0:
            result[src] = produced
    return result


# --------------------------------------------------------------------------
# Format-specific extraction
# --------------------------------------------------------------------------


def _iter_docx_blocks(document):
    """Yield paragraphs and tables in true document order.

    python-docx exposes .paragraphs and .tables as separate flat lists, which
    loses their interleaving; walking the body children preserves it.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield Paragraph(child, document)
        elif tag == "tbl":
            yield Table(child, document)


def _table_to_text(table) -> str:
    """Render a table as tab-separated rows, preserving its content."""
    lines = []
    for row in table.rows:
        cells = [" ".join(c.text.split()) for c in row.cells]
        # Collapse the horizontally-merged repeats python-docx reports per cell.
        deduped = [c for i, c in enumerate(cells) if i == 0 or c != cells[i - 1]]
        if any(deduped):
            lines.append("\t".join(deduped))
    return "\n".join(lines)


def extract_docx(path: Path) -> list[Page]:
    """Extract a .docx. Word has no reliable page model, so this yields one page
    with number=None (see ExtractedDoc.page_number_note)."""
    import docx
    from docx.table import Table

    try:
        document = docx.Document(str(path))
    except Exception as exc:  # python-docx raises assorted types on damage
        raise ExtractionError(f"docx parse failed: {type(exc).__name__}: {exc}") from exc

    parts: list[str] = []
    for block in _iter_docx_blocks(document):
        if isinstance(block, Table):
            rendered = _table_to_text(block)
            if rendered.strip():
                parts.append(rendered)
        else:
            text = block.text
            if text and text.strip():
                parts.append(text.strip())
    return [Page(number=None, text="\n".join(parts))]


def extract_pdf(path: Path) -> list[Page]:
    """Extract a PDF page by page. Pages with no usable text layer are kept and
    marked needs_ocr rather than dropped."""
    import pymupdf

    try:
        document = pymupdf.open(str(path))
    except Exception as exc:
        raise ExtractionError(f"pdf open failed: {type(exc).__name__}: {exc}") from exc

    pages: list[Page] = []
    with document:
        for index, page in enumerate(document, start=1):
            try:
                text = page.get_text("text") or ""
            except Exception as exc:
                pages.append(Page(number=index, text="", needs_ocr=True))
                continue
            stripped = text.strip()
            pages.append(
                Page(
                    number=index,
                    text=stripped,
                    needs_ocr=len(stripped) < config.MIN_PAGE_CHARS,
                )
            )
    return pages


# --------------------------------------------------------------------------
# Text normalization
# --------------------------------------------------------------------------


def _is_lower_tr(ch: str) -> bool:
    """Turkish-aware lowercase test (str.islower() mishandles ı/İ round-trips)."""
    return ch.isalpha() and tr_lower(ch) == ch


def fix_hyphenation(text: str) -> str:
    """Rejoin words split across a line break, keeping genuine hyphens.

    "yönet-\nmelik" -> "yönetmelik", but "ön-\nKoşul" is left alone because an
    uppercase continuation signals a real compound rather than a line-break
    split. A hyphen not at a line break (ön-koşul) is never touched.
    """
    out: list[str] = []
    i = 0
    n = len(text)
    while i < n:
        ch = text[i]
        if ch == "-":
            j = i + 1
            # Consume the line break and any indentation that follows it.
            saw_newline = False
            while j < n and text[j] in " \t\r":
                j += 1
            if j < n and text[j] == "\n":
                saw_newline = True
                j += 1
                while j < n and text[j] in " \t\r":
                    j += 1
            if saw_newline and j < n and _is_lower_tr(text[j]) and i > 0 and text[i - 1].isalpha():
                # Drop the hyphen and the break, splicing the halves together.
                i = j
                continue
        out.append(ch)
        i += 1
    return "".join(out)


def _looks_like_page_number(line: str) -> bool:
    stripped = line.strip().strip("-–—.[]() ")
    if not stripped or len(stripped) > 12:
        return False
    if stripped.isdigit():
        return True
    low = tr_lower(stripped)
    # "Sayfa 3", "3/12"
    if low.startswith("sayfa") and any(c.isdigit() for c in low):
        return True
    parts = stripped.split("/")
    return len(parts) == 2 and all(p.strip().isdigit() for p in parts)


def strip_repeating_lines(pages: list[Page], min_repeats: int | None = None) -> list[Page]:
    """Remove running headers/footers and standalone page numbers.

    A short line appearing near the top or bottom of most pages is treated as
    furniture. Requires several pages to judge, so single-page documents only
    get their standalone page numbers removed.
    """
    page_count = len(pages)
    threshold = min_repeats if min_repeats is not None else max(2, int(page_count * 0.6))

    repeated: set[str] = set()
    if page_count >= 3:
        counts: dict[str, int] = {}
        for page in pages:
            lines = [ln.strip() for ln in page.text.splitlines() if ln.strip()]
            edge = lines[:2] + lines[-2:]
            for line in set(edge):
                if len(line) <= 90:
                    counts[line] = counts.get(line, 0) + 1
        repeated = {line for line, n in counts.items() if n >= threshold}

    cleaned: list[Page] = []
    for page in pages:
        kept = [
            ln
            for ln in page.text.splitlines()
            if ln.strip() not in repeated and not _looks_like_page_number(ln)
        ]
        cleaned.append(
            Page(number=page.number, text="\n".join(kept).strip(), needs_ocr=page.needs_ocr)
        )
    return cleaned


def normalize_text(text: str) -> str:
    """NFC-normalize, fix hyphenation and collapse excess blank lines."""
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ").replace("\u200b", "")
    text = fix_hyphenation(text)
    lines = [ln.rstrip() for ln in text.split("\n")]
    out: list[str] = []
    blanks = 0
    for line in lines:
        if line.strip():
            blanks = 0
            out.append(line)
        else:
            blanks += 1
            if blanks <= 1:
                out.append("")
    return "\n".join(out).strip()


# --------------------------------------------------------------------------
# Quality checks
# --------------------------------------------------------------------------


def _longest_nonspace_run(text: str) -> int:
    best = current = 0
    for ch in text:
        if ch.isspace():
            current = 0
        else:
            current += 1
            best = max(best, current)
    return best


# Sequences produced by decoding UTF-8 Turkish as cp1252/latin-1.
_MOJIBAKE_MARKERS = ("Ã¼", "Ã¶", "Ã§", "Ãı", "Å", "Ä°", "Ä±", "Ã\x9c", "ÃÂ")


def quality_flags(doc: ExtractedDoc) -> list[str]:
    """Return human-readable quality flags. Flagging never removes a document."""
    flags: list[str] = []
    text = doc.text

    ocr_pages = [p.number for p in doc.pages if p.needs_ocr]

    if not text.strip():
        # Report the *cause* alongside the symptom: an empty PDF whose pages all
        # lack a text layer is a scan needing OCR, which is far more actionable
        # than a bare "empty".
        if ocr_pages and len(ocr_pages) == len(doc.pages):
            flags.append(
                f"needs-ocr: no text layer on any of {len(doc.pages)} page(s) - "
                "likely a scanned document"
            )
        flags.append("empty: no text extracted")
        return flags
    if len(text) < config.MIN_DOC_CHARS:
        flags.append(f"near-empty: only {len(text)} chars extracted")

    replacements = text.count("\ufffd")
    if replacements and replacements / max(len(text), 1) > config.MAX_REPLACEMENT_RATIO:
        flags.append(f"replacement-chars: {replacements} U+FFFD present (bad decode)")

    hits = [m for m in _MOJIBAKE_MARKERS if m in text]
    if hits:
        flags.append(f"mojibake: found {hits[:3]} (UTF-8 text decoded as cp1252?)")

    turkish = sum(text.count(c) for c in TURKISH_CHARS)
    if turkish < config.MIN_TURKISH_CHAR_COUNT:
        flags.append("no-turkish-chars: 0 of ıİşğüöç in a Turkish document (mis-decoded?)")

    run = _longest_nonspace_run(text)
    if run > config.MAX_GIBBERISH_RUN:
        flags.append(f"gibberish-run: {run} chars with no whitespace")

    if ocr_pages:
        shown = ", ".join(str(p) for p in ocr_pages[:8])
        more = f" (+{len(ocr_pages) - 8} more)" if len(ocr_pages) > 8 else ""
        flags.append(f"needs-ocr: {len(ocr_pages)} page(s) with no text layer: {shown}{more}")

    empty_pages = [p.number for p in doc.pages if p.number and not p.text.strip() and not p.needs_ocr]
    if empty_pages:
        flags.append(f"empty-pages: {len(empty_pages)} page(s) extracted blank")

    return flags


# --------------------------------------------------------------------------
# Orchestration
# --------------------------------------------------------------------------

# Non-prose formats present in this corpus. Recorded with a reason so they show
# up in the report instead of vanishing.
_SKIP_REASONS = {
    "xlsx": "spreadsheet (.xlsx) - tabular data, not regulation prose",
    "xls": "spreadsheet (.xls) - tabular data, not regulation prose",
    "zip": "archive (.zip) - contains nested files; not unpacked by this pass",
    "pptx": "presentation (.pptx) - not regulation prose",
    "ppt": "presentation (.ppt) - not regulation prose",
    "rtf": "rich text (.rtf) - no extractor wired up",
    "html": "html - no extractor wired up",
    "empty": "file is 0 bytes",
}


def iter_corpus_files(root: Path) -> list[Path]:
    """All files under root, recursively, in a stable sorted order."""
    return sorted(p for p in root.rglob("*") if p.is_file())


def extract_file(path: Path, converted: Path | None = None) -> ExtractedDoc:
    """Extract one file that is already known to be an extractable type."""
    file_hash = sha256_file(path)
    detected = detect_type(path)
    doc = ExtractedDoc(
        path=path,
        original_filename=path.name,
        doc_id=make_doc_id(path, file_hash),
        file_sha256=file_hash,
        detected_type=detected,
    )

    if detected == "pdf":
        pages = extract_pdf(path)
    elif detected == "docx":
        pages = extract_docx(path)
        doc.page_number_note = "docx has no fixed page model; page numbers unavailable"
    elif detected == "doc":
        if converted is None:
            raise ExtractionError("no LibreOffice conversion available for .doc")
        pages = extract_docx(converted)
        doc.page_number_note = (
            "converted from .doc via LibreOffice; page numbers unavailable"
        )
    else:
        raise ExtractionError(f"unsupported type {detected!r}")

    pages = strip_repeating_lines(pages)
    doc.pages = [
        Page(number=p.number, text=normalize_text(p.text), needs_ocr=p.needs_ocr) for p in pages
    ]
    doc.flags = quality_flags(doc)
    return doc


@dataclass
class ExtractionRun:
    """The outcome of extracting a whole corpus: what worked, what didn't, and why."""

    docs: list[ExtractedDoc] = field(default_factory=list)
    skipped: list[SkippedDoc] = field(default_factory=list)
    failed: list[SkippedDoc] = field(default_factory=list)
    duplicates: list[tuple[Path, str]] = field(default_factory=list)
    type_counts: dict[str, int] = field(default_factory=dict)


def extract_corpus(
    root: Path,
    limit: int | None = None,
    verbose: bool = True,
    paths: list[Path] | None = None,
) -> ExtractionRun:
    """Extract every extractable document under `root`.

    `paths` extracts an explicit subset instead of scanning `root`, so a partial
    reprocess reuses this function's .doc batching, duplicate detection and
    error handling rather than reimplementing them. Note that duplicate
    detection is then scoped to the subset: two files with identical content
    both get extracted if only one is in `paths`, which is correct -- the
    caller asked for those specific documents.

    Fails fast (LibreOfficeUnavailable) when .doc files are present but
    LibreOffice is missing, rather than skipping a quarter of the corpus.
    """
    run = ExtractionRun()
    files = list(paths) if paths is not None else iter_corpus_files(root)
    if limit:
        files = files[:limit]

    # Classify everything first so the .doc conversion can be batched.
    classified: list[tuple[Path, str]] = []
    for path in files:
        try:
            detected = detect_type(path)
        except ExtractionError as exc:
            run.failed.append(SkippedDoc(path, str(exc)))
            continue
        run.type_counts[detected] = run.type_counts.get(detected, 0) + 1
        classified.append((path, detected))

    doc_paths = [p for p, t in classified if t == "doc"]
    conversions: dict[Path, Path] = {}
    tmpdir: tempfile.TemporaryDirectory | None = None

    if doc_paths:
        soffice = find_libreoffice()
        if soffice is None:
            raise LibreOfficeUnavailable(
                f"{len(doc_paths)} .doc file(s) need LibreOffice but 'soffice' was not found.\n"
                "Install LibreOffice or set LIBREOFFICE_PATH in src/config.py.\n"
                "Refusing to continue: skipping these would silently drop "
                f"{len(doc_paths)} of {len(classified)} corpus files."
            )
        if verbose:
            print(f"LibreOffice: {soffice}")
            print(f"Converting {len(doc_paths)} .doc file(s) to .docx ...", flush=True)

        tmpdir = tempfile.TemporaryDirectory(prefix="epdk-doc-")
        outdir = Path(tmpdir.name)
        size = max(1, config.LIBREOFFICE_BATCH_SIZE)
        for start in range(0, len(doc_paths), size):
            batch = doc_paths[start : start + size]
            conversions.update(convert_doc_batch(batch, soffice, outdir))
            if verbose:
                done = min(start + size, len(doc_paths))
                print(f"  converted {len(conversions)}/{done} (of {len(doc_paths)})", flush=True)

    try:
        seen_hashes: dict[str, str] = {}
        for path, detected in classified:
            if detected not in EXTRACTABLE:
                reason = _SKIP_REASONS.get(detected, f"unsupported type: {detected}")
                run.skipped.append(SkippedDoc(path, reason))
                continue

            if detected == "doc" and path not in conversions:
                run.failed.append(
                    SkippedDoc(path, "LibreOffice conversion produced no output (corrupt .doc?)")
                )
                continue

            try:
                doc = extract_file(path, conversions.get(path))
            except ExtractionError as exc:
                run.failed.append(SkippedDoc(path, str(exc)))
                continue
            except Exception as exc:  # noqa: BLE001 - report, never crash the run
                run.failed.append(SkippedDoc(path, f"unexpected {type(exc).__name__}: {exc}"))
                continue

            if doc.file_sha256 in seen_hashes:
                run.duplicates.append((path, seen_hashes[doc.file_sha256]))
                continue
            seen_hashes[doc.file_sha256] = doc.doc_id
            run.docs.append(doc)
    finally:
        if tmpdir is not None:
            tmpdir.cleanup()

    return run


# --------------------------------------------------------------------------
# CLI
# --------------------------------------------------------------------------


def _force_utf8_stdout() -> None:
    """Windows consoles default to cp1254 here, which raises on Turkish output."""
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except (AttributeError, OSError):
            pass


def _rel(path: Path, root: Path) -> str:
    try:
        return str(path.relative_to(root))
    except ValueError:
        return str(path)


def main(argv: list[str] | None = None) -> int:
    """CLI entry point: extract the corpus and print an extraction report."""
    import argparse

    from .chunk import chunk_document
    from .titles import extract_title

    _force_utf8_stdout()
    parser = argparse.ArgumentParser(prog="python -m src.extract")
    parser.add_argument("--root", default="data", help="corpus root (default: data)")
    parser.add_argument("--limit", type=int, default=None, help="process at most N files")
    parser.add_argument("--show-flagged", type=int, default=40, help="max flagged docs to list")
    parser.add_argument("--sample", type=int, default=0, help="print N extracted samples")
    args = parser.parse_args(argv)

    root = Path(args.root)
    if not root.exists():
        print(f"error: corpus root {root} does not exist")
        return 2

    print(f"Scanning {root.resolve()} ...")
    try:
        run = extract_corpus(root, limit=args.limit)
    except LibreOfficeUnavailable as exc:
        print(f"\nFATAL: {exc}")
        return 3

    # Titles and chunks
    titles_found = 0
    title_missing: list[ExtractedDoc] = []
    total_chunks = 0
    chunks_with_pages = 0
    strategy_counts: dict[str, int] = {}

    for doc in run.docs:
        info = extract_title(doc)
        if info.title:
            titles_found += 1
        else:
            title_missing.append(doc)
        doc.flags.extend(info.flags)

        chunks = chunk_document(doc, info)
        total_chunks += len(chunks)
        for c in chunks:
            strategy_counts[c.strategy] = strategy_counts.get(c.strategy, 0) + 1
            if c.page_start is not None:
                chunks_with_pages += 1

    n_docs = len(run.docs)
    print("\n" + "=" * 72)
    print("EXTRACTION REPORT")
    print("=" * 72)

    print(f"\nFiles seen              : {sum(run.type_counts.values())}")
    print(f"Documents extracted     : {n_docs}")
    print(f"Chunks produced         : {total_chunks}")
    print(f"Titles found            : {titles_found}")
    print(f"Titles NOT found        : {len(title_missing)}")

    print("\n-- File types --")
    ok_by_type: dict[str, int] = {}
    for doc in run.docs:
        ok_by_type[doc.detected_type] = ok_by_type.get(doc.detected_type, 0) + 1
    fail_by_type: dict[str, int] = {}
    for item in run.failed:
        try:
            t = detect_type(item.path)
        except ExtractionError:
            t = "unreadable"
        fail_by_type[t] = fail_by_type.get(t, 0) + 1
    for t in sorted(run.type_counts, key=lambda k: -run.type_counts[k]):
        mark = "extractable" if t in EXTRACTABLE else "skipped"
        print(
            f"  {t:10} seen={run.type_counts[t]:4}  ok={ok_by_type.get(t, 0):4}  "
            f"failed={fail_by_type.get(t, 0):3}  [{mark}]"
        )

    print("\n-- Chunking strategy --")
    for name in sorted(strategy_counts, key=lambda k: -strategy_counts[k]):
        print(f"  {name:22} {strategy_counts[name]}")
    pages_note = f"{chunks_with_pages}/{total_chunks} chunks carry page numbers"
    print(f"\n  {pages_note} (PDF only; .doc/.docx have no page model)")

    if run.duplicates:
        print(f"\n-- Duplicates skipped (identical SHA-256): {len(run.duplicates)} --")
        for path, first_id in run.duplicates[:10]:
            print(f"  {_rel(path, root)}\n      same content as {first_id}")
        if len(run.duplicates) > 10:
            print(f"  ... and {len(run.duplicates) - 10} more")

    if run.skipped:
        print(f"\n-- Skipped, not an extractable document type: {len(run.skipped)} --")
        by_reason: dict[str, list[Path]] = {}
        for item in run.skipped:
            by_reason.setdefault(item.reason, []).append(item.path)
        for reason, paths in sorted(by_reason.items(), key=lambda kv: -len(kv[1])):
            print(f"  [{len(paths):3}] {reason}")
            for p in paths[:3]:
                print(f"          {_rel(p, root)}")
            if len(paths) > 3:
                print(f"          ... and {len(paths) - 3} more")

    if run.failed:
        print(f"\n-- FAILED to extract: {len(run.failed)} --")
        for item in run.failed:
            print(f"  {_rel(item.path, root)}\n      reason: {item.reason}")

    flagged = [d for d in run.docs if d.flags]
    print(f"\n-- Quality flags: {len(flagged)} document(s) flagged for review --")
    if flagged:
        counts: dict[str, int] = {}
        for doc in flagged:
            for flag in doc.flags:
                counts[flag.split(":")[0]] = counts.get(flag.split(":")[0], 0) + 1
        print("   by check:")
        for name in sorted(counts, key=lambda k: -counts[k]):
            print(f"     {name:22} {counts[name]}")
        print()
        for doc in flagged[: args.show_flagged]:
            print(f"  {doc.doc_id}  [{doc.detected_type}]")
            print(f"      file: {_rel(doc.path, root)}")
            for flag in doc.flags:
                print(f"      FLAG: {flag}")
        if len(flagged) > args.show_flagged:
            print(f"  ... and {len(flagged) - args.show_flagged} more (use --show-flagged)")

    if title_missing:
        print(f"\n-- Title not determined: {len(title_missing)} --")
        for doc in title_missing[:15]:
            print(f"  {doc.doc_id}  {_rel(doc.path, root)}")
        if len(title_missing) > 15:
            print(f"  ... and {len(title_missing) - 15} more")

    if args.sample:
        print("\n-- Extracted samples (Turkish integrity check) --")
        for doc in run.docs[: args.sample]:
            info = extract_title(doc)
            snippet = " ".join(doc.text.split())[:220]
            turkish = sum(doc.text.count(c) for c in TURKISH_CHARS)
            print(f"\n  {doc.doc_id} [{doc.detected_type}]")
            print(f"    title : {info.title}")
            print(f"    type  : {info.mevzuat_type}   number: {info.number}")
            print(f"    tr chars: {turkish}")
            print(f"    text  : {snippet}")

    print()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
